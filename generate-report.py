#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерация пояснительной записки к курсовой работе CarRentalApp"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ============ Стили ============
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(1.25)

# Настройка полей
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

def add_heading_styled(text, level=1):
    """Добавить заголовок с Times New Roman"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = bold
    run.italic = italic
    if align:
        p.paragraph_format.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.5)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    doc.add_paragraph()

# ============ ТИТУЛЬНЫЙ ЛИСТ ============
for _ in range(6):
    doc.add_paragraph()

add_para("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Федеральное государственное автономное образовательное учреждение\nвысшего образования", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("\"СЕВЕРО-КАВКАЗСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ\"", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para("Институт цифрового развития", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Кафедра информационных систем и технологий", align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

add_para("КУРСОВАЯ РАБОТА", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("по дисциплине «Программная инженерия»", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("на тему: «Информационная система проката автомобилей»", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("Выполнил:\nстудент 4 курса, группы ПИЖ-б-о-21-1\nДжабраилов Б.С-Х.")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("\nПроверил:\nдоцент кафедры ИСТ\nКириллов И.В.")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()
add_para("Ставрополь, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============ АННОТАЦИЯ ============
add_heading_styled("АННОТАЦИЯ", 1)
add_para("Данная курсовая работа посвящена разработке информационной системы проката автомобилей «CarRentalApp» в рамках траектории «Мобильная разработка». Целью работы является создание полнофункционального мобильного приложения для бронирования автомобилей, включающего серверную часть на базе Spring Boot и клиентское приложение на Android с использованием Jetpack Compose.")
add_para("В ходе работы были решены следующие задачи: проведён анализ предметной области, спроектирована архитектура системы с использованием паттерна PCMEF, разработана база данных (PostgreSQL) с 6 таблицами, реализован REST API (20+ endpoint'ов) с JWT-аутентификацией, создано Android-приложение с 7 экранами на Jetpack Compose, обеспечена работа в офлайн-режиме через Room, проведено тестирование (12 тестовых классов, покрытие >40%), настроена Docker-контейнеризация и CI/CD через GitHub Actions.")
add_para("Итоговая оценка по COCOMO: трудоёмкость 8.6 человеко-месяца, срок разработки ~6 месяцев. Общий размер кода составляет ~9 000 строк.")

doc.add_page_break()

# ============ СОДЕРЖАНИЕ ============
add_heading_styled("СОДЕРЖАНИЕ", 1)
toc_items = [
    ("ВВЕДЕНИЕ", 3),
    ("1 АНАЛИТИЧЕСКАЯ ЧАСТЬ", 5),
    ("1.1 Технико-экономическая характеристика объекта", 5),
    ("1.2 Архитектура системы (PCMEF)", 7),
    ("1.3 Обоснование проектных решений", 9),
    ("2 ПРОЕКТНАЯ ЧАСТЬ", 11),
    ("2.1 Программная реализация системы", 11),
    ("2.2 Тестирование", 15),
    ("2.3 Руководство пользователя", 17),
    ("3 ЭКОНОМИЧЕСКАЯ ЧАСТЬ", 20),
    ("3.1 WBS — иерархическая структура работ", 20),
    ("3.2 Gantt-диаграмма (календарный план)", 21),
    ("3.3 COCOMO — оценка трудоёмкости", 22),
    ("ЗАКЛЮЧЕНИЕ", 24),
    ("СПИСОК ЛИТЕРАТУРЫ", 25),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{item}{'.' * (60 - len(item))}{page}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)

doc.add_page_break()

# ============ ВВЕДЕНИЕ ============
add_heading_styled("ВВЕДЕНИЕ", 1)
add_para("Современный рынок аренды автомобилей характеризуется высокой динамикой и растущими требованиями к удобству и скорости обслуживания клиентов. Традиционные процессы бронирования, предполагающие телефонные звонки и личное посещение офисов, уступают место цифровым решениям, позволяющим арендовать автомобиль в несколько кликов со смартфона.")
add_para("Актуальность разработки обусловлена растущим спросом на мобильные сервисы в сфере car sharing и краткосрочной аренды автомобилей, а также необходимостью автоматизации процессов управления автопарком для компаний-арендодателей.")
add_para("Целью данной курсовой работы является разработка информационной системы проката автомобилей «CarRentalApp», обеспечивающей удобный интерфейс для клиентов и инструменты администрирования для управления автопарком.")
add_para("Для достижения поставленной цели необходимо решить следующие задачи:", bold=True)
add_bullet("Провести анализ предметной области и существующих решений")
add_bullet("Спроектировать архитектуру системы на основе паттерна PCMEF")
add_bullet("Разработать базу данных (PostgreSQL) с нормализацией до 3НФ")
add_bullet("Реализовать серверную часть на Java Spring Boot с REST API")
add_bullet("Разработать Android-клиент на Kotlin с Jetpack Compose")
add_bullet("Обеспечить безопасность данных (JWT, BCrypt, ролевая модель)")
add_bullet("Реализовать офлайн-режим работы через Room")
add_bullet("Провести тестирование (JUnit, Mockito) с покрытием >40%")
add_bullet("Настроить Docker-контейнеризацию и CI/CD")
add_para("Работа выполнена в рамках траектории «Мобильная разработка» с использованием следующего стека технологий: Java 17, Spring Boot 3.2, PostgreSQL, Kotlin, Jetpack Compose, Retrofit, Room, Docker.")

doc.add_page_break()

# ============ ГЛАВА 1 ============
add_heading_styled("1 АНАЛИТИЧЕСКАЯ ЧАСТЬ", 1)

add_heading_styled("1.1 Технико-экономическая характеристика объекта", 2)
add_para("Объектом автоматизации является процесс предоставления автомобилей в аренду. В традиционной схеме клиент обращается в офис компании, заключает договор аренды, вносит залог и получает автомобиль. Данный процесс требует личного присутствия, сопряжён с бумажным документооборотом и затрудняет оперативный поиск доступных автомобилей.")
add_para("Разрабатываемая информационная система призвана автоматизировать следующие бизнес-процессы:")
add_bullet("Поиск автомобилей по фильтрам (класс, даты, цена)")
add_bullet("Бронирование автомобиля с проверкой доступности")
add_bullet("Оплата аренды (в учебной реализации — симуляция)")
add_bullet("Выдача и возврат автомобиля менеджером")
add_bullet("Управление автопарком (добавление, редактирование, удаление ТС)")
add_bullet("Верификация клиентов администратором")
add_para("Основные роли пользователей системы:", bold=True)
add_table(
    ["Роль", "Описание", "Права"],
    [
        ["CLIENT", "Клиент (арендатор)", "Поиск, бронирование, просмотр истории"],
        ["MANAGER", "Менеджер проката", "Выдача/приём авто, просмотр броней"],
        ["ADMIN", "Администратор", "Полный доступ: управление автопарком, клиентами"],
    ]
)

add_heading_styled("1.1.1 Организационная структура", 3)
add_para("Взаимодействие пользователей с системой организовано через три уровня доступа. Клиент взаимодействует с мобильным приложением, менеджер использует панель управления для операционных задач (выдача и приём автомобилей), администратор управляет всей системой.")

add_heading_styled("1.2 Архитектура системы (PCMEF)", 2)
add_heading_styled("1.2.1 Обоснование выбора архитектуры", 3)
add_para("В качестве архитектурного паттерна выбран PCMEF (Presentation — Control — Mediator — Entity — Foundation), рекомендованный методическими указаниями. Данный паттерн обеспечивает строгую Layer Isolation: каждый слой зависит только от нижележащего, что улучшает testability и maintainability системы.")

add_heading_styled("1.2.2 Схема архитектуры PCMEF", 3)
add_para("Архитектура системы включает два крупных компонента — серверный (Backend) и клиентский (Android):")
add_para("Серверная часть (Java 17, Spring Boot 3.2):", bold=True)
add_bullet("Control — REST-контроллеры (AuthController, CarController, BookingController, AdminController)")
add_bullet("Mediator — сервисы бизнес-логики (AuthService, CarService, ReservationService, BookingFacade)")
add_bullet("Entity — JPA-сущности (User, Car, Reservation, RentalAgreement, Payment, ClientProfile)")
add_bullet("Foundation — JPA-репозитории (Spring Data JPA)")
add_para("Клиентская часть (Kotlin, Jetpack Compose):", bold=True)
add_bullet("Presentation — экраны (Login, Catalog, Detail, Booking, Profile, Admin)")
add_bullet("StateManagement — ViewModel + StateFlow (MVVM)")
add_bullet("ApiClient — Retrofit + OkHttp + JWT Interceptor")
add_bullet("LocalCache — Room (офлайн-кэш) + EncryptedSharedPreferences")

add_heading_styled("1.2.3 Реализованные паттерны GoF", 3)
add_table(
    ["Паттерн", "Описание", "Реализация"],
    [
        ["State", "Жизненный цикл бронирования", "ReservationState → Pending/Confirmed/Active/CancelledState"],
        ["Strategy", "Ценообразование", "PricingStrategy → Standard/WeekendPricingStrategy"],
        ["Facade", "Сложный процесс бронирования", "BookingFacade.processCompleteBooking()"],
        ["Singleton", "Единый экземпляр", "ApiClient, AppDatabase, TokenStorage"],
    ]
)

add_heading_styled("1.3 Обоснование проектных решений", 2)
add_heading_styled("1.3.1 Техническое обеспечение", 3)
add_table(
    ["Компонент", "Технология", "Обоснование"],
    [
        ["Backend", "Java 17 + Spring Boot 3.2", "Стабильность, экосистема, соответствие требованиям"],
        ["БД", "PostgreSQL 15", "Надёжность, поддержка JSON, Flyway миграции"],
        ["ORM", "Spring Data JPA / Hibernate", "Декларативные репозитории, кэш 2-го уровня"],
        ["API", "REST + OpenAPI/Swagger", "Стандарт индустрии, автодокументация"],
        ["Auth", "JWT + BCrypt + Spring Security", "Stateless, безопасное хранение паролей"],
        ["Android", "Kotlin + Jetpack Compose", "Современная declarative UI, type-safe"],
        ["Сеть", "Retrofit 2 + OkHttp", "Асинхронные запросы, интерцепторы"],
        ["Локальное хранение", "Room", "Type-safe SQLite, корутины, миграции"],
    ]
)

add_heading_styled("1.3.2 Информационное обеспечение", 3)
add_para("База данных спроектирована в 3-й нормальной форме. Она включает 6 таблиц:")
add_bullet("users — учётные записи пользователей (email, password_hash, role)")
add_bullet("client_profiles — профили клиентов (верификация, паспортные данные)")
add_bullet("cars — автомобили (VIN, госномер, модель, статус, тариф)")
add_bullet("reservations — бронирования (клиент, авто, даты, статус, сумма)")
add_bullet("rental_agreements — договоры аренды (пробег, топливо, активность)")
add_bullet("payments — платежи (транзакция, тип, статус)")
add_para("Применены ограничения: NOT NULL, UNIQUE (email, VIN, госномер), CHECK (статусы), FOREIGN KEY с каскадным удалением. Миграции управляются через Flyway (V1, V2, V3).")

doc.add_page_break()

# ============ ГЛАВА 2 ============
add_heading_styled("2 ПРОЕКТНАЯ ЧАСТЬ", 1)

add_heading_styled("2.1 Программная реализация системы", 2)

add_heading_styled("2.1.1 Структура проекта", 3)
add_para("Проект организован в соответствии с архитектурой PCMEF и содержит следующие модули:")

add_para("Backend (серверная часть):", bold=True)
add_bullet("control — 4 REST-контроллера (Auth, Car, Booking, Admin)")
add_bullet("mediator — 6 сервисов (Auth, Car, Reservation, Payment, Notification, BookingFacade)")
add_bullet("entity — 6 JPA-сущностей + 4 enum + State-паттерн")
add_bullet("foundation — 6 JPA-репозиториев")
add_bullet("security — JwtAuthFilter, JwtTokenProvider, SecurityConfig")
add_bullet("exception — GlobalExceptionHandler + кастомные исключения")
add_bullet("scheduler — планировщик автоотмены броней")

add_para("Android (клиентская часть):", bold=True)
add_bullet("presentation — 7 экранов на Jetpack Compose + навигация")
add_bullet("statemanagement — 5 ViewModel'ей с StateFlow")
add_bullet("apiclient — Retrofit-сервисы + JWT-интерцептор")
add_bullet("localcache — Room (AppDatabase, DAO, TokenStorage)")
add_bullet("model — DTO для запросов/ответов")
add_bullet("ui/theme — Material 3 тема (светлая/тёмная)")

add_heading_styled("2.1.2 REST API", 3)
add_para("Система предоставляет 20+ REST endpoint'ов, разделённых на группы:")

add_para("Публичные (Auth):", bold=True)
add_table(["Метод", "Endpoint", "Описание"], [
    ["POST", "/api/v1/auth/login", "Аутентификация, возвращает JWT"],
    ["POST", "/api/v1/auth/register", "Регистрация нового пользователя"],
])

add_para("Каталог (Cars):", bold=True)
add_table(["Метод", "Endpoint", "Роль"], [
    ["GET", "/api/v1/cars", "Любая"],
    ["GET", "/api/v1/cars/{id}", "Любая"],
    ["GET", "/api/v1/cars/{id}/busy", "Любая"],
    ["PUT", "/api/v1/cars/{id}/status", "MANAGER/ADMIN"],
])

add_para("Бронирования (Bookings):", bold=True)
add_table(["Метод", "Endpoint", "Роль"], [
    ["POST", "/api/v1/bookings", "CLIENT"],
    ["GET", "/api/v1/bookings/my", "CLIENT"],
    ["GET", "/api/v1/bookings/{id}", "CLIENT"],
    ["POST", "/api/v1/bookings/{id}/cancel", "CLIENT"],
])

add_para("Администрирование (Admin):", bold=True)
add_table(["Метод", "Endpoint", "Роль"], [
    ["GET", "/api/v1/admin/cars", "MANAGER/ADMIN"],
    ["POST/PUT/DELETE", "/api/v1/admin/cars[/{id}]", "ADMIN"],
    ["GET", "/api/v1/admin/bookings", "MANAGER/ADMIN"],
    ["POST", "/api/v1/admin/bookings/{id}/handover", "MANAGER"],
    ["POST", "/api/v1/admin/bookings/{id}/return", "MANAGER"],
    ["GET", "/api/v1/admin/clients/unverified", "MANAGER/ADMIN"],
    ["PUT", "/api/v1/admin/clients/{userId}/verify", "ADMIN"],
])

add_heading_styled("2.1.3 Реализация State-паттерна", 3)
add_para("Жизненный цикл бронирования реализован через паттерн State. Каждое состояние представлено отдельным классом:")
add_bullet("PendingState — бронирование создано, ожидает оплаты (автоотмена через 15 минут)")
add_bullet("ConfirmedState — оплата подтверждена, ожидает выдачи автомобиля")
add_bullet("ActiveState — автомобиль выдан клиенту (активная аренда)")
add_bullet("CancelledState — бронирование отменено (клиентом или автоотмена)")
add_bullet("COMPLETED — аренда завершена, автомобиль возвращён")

add_heading_styled("2.2 Тестирование", 2)
add_para("Тестирование проведено на двух уровнях:")

add_para("Модульное тестирование (JUnit 5 + Mockito):", bold=True)
add_bullet("AuthServiceTest — 6 тестов (логин, регистрация, валидация)")
add_bullet("CarServiceTest — 5 тестов (поиск, фильтрация, получение по ID)")
add_bullet("CarServiceUpdateTest — 3 теста (обновление, ошибки)")
add_bullet("CarServiceExtendedTest — 3 теста (статус, создание)")
add_bullet("ReservationServiceTest — 4 теста (бронирование, исключения)")
add_bullet("ReservationServiceExtendedTest — 5 тестов (отмена, занятые даты)")
add_bullet("ReservationServiceHandoverReturnTest — 8 тестов (выдача/возврат)")
add_bullet("ReservationStateTest — тесты State-паттерна")
add_bullet("PricingStrategyTest — тесты ценообразования")
add_bullet("GlobalExceptionHandlerTest — 7 тестов (обработка ошибок)")

add_para("Интеграционное тестирование (Spring Boot Test + MockMvc + H2):", bold=True)
add_bullet("CarControllerIntegrationTest — 5 тестов (авторизация, фильтрация, поиск)")
add_bullet("SecurityConfigTest — 3 теста (публичные endpoint'ы)")
add_para("Общее покрытие кода тестами превышает 40% (JaCoCo).")

add_heading_styled("2.3 Руководство пользователя", 2)

add_heading_styled("2.3.1 Клиентское приложение", 3)
add_para("Мобильное приложение включает 7 экранов:")
add_bullet("Экран входа — аутентификация по email/паролю")
add_bullet("Экран регистрации — создание новой учётной записи")
add_bullet("Каталог — список доступных автомобилей с фильтрацией по классу, датам и поиском по модели")
add_bullet("Детали авто — подробная информация, фото, занятые даты, кнопка бронирования")
add_bullet("Бронирование — выбор дат, расчёт стоимости, подтверждение")
add_bullet("Профиль — список бронирований, отмена, статус верификации")
add_bullet("Админ-панель — управление автопарком, бронированиями, клиентами")

add_heading_styled("2.3.2 Панель администратора", 3)
add_para("Панель администратора доступна пользователям с ролями ADMIN и MANAGER. Она содержит три вкладки:")
add_bullet("Автопарк — просмотр, добавление, редактирование, удаление автомобилей, смена статуса")
add_bullet("Бронирования — просмотр всех броней, отмена, выдача и приём автомобиля")
add_bullet("Клиенты — верификация документов новых клиентов")

doc.add_page_break()

# ============ ГЛАВА 3 ============
add_heading_styled("3 ЭКОНОМИЧЕСКАЯ ЧАСТЬ", 1)

add_heading_styled("3.1 WBS — иерархическая структура работ", 2)
add_para("Проект включает 6 основных блоков работ:")
add_bullet("1. Управление проектом — планирование, контроль версий, документирование")
add_bullet("2. Анализ и проектирование — сбор требований, BUC, IDEF0, PCMEF, ER, DDL")
add_bullet("3. Разработка серверной части — Spring Boot, JPA, JWT, REST API, тесты")
add_bullet("4. Разработка клиентской части — Jetpack Compose, MVVM, Retrofit, Room")
add_bullet("5. Интеграция и развёртывание — Docker, CI/CD, интеграционное тестирование")
add_bullet("6. Документирование — пояснительная записка, презентация")

add_heading_styled("3.2 Gantt-диаграмма (календарный план)", 2)
add_para("Календарный план включает 21 задачу на 10 недель (2.5 месяца). Ключевые вехи:")
add_table(
    ["ID", "Задача", "Длит.", "Недели"],
    [
        ["1-2", "Управление + сбор требований", "1 нед", "Н1"],
        ["3-4", "Анализ и проектирование", "2 нед", "Н2-Н3"],
        ["5-6", "Проектирование БД и API", "1 нед", "Н3"],
        ["7-10", "Разработка серверной части", "3 нед", "Н4-Н6"],
        ["11-14", "Разработка Android", "3.5 нед", "Н5-Н8"],
        ["15-16", "Офлайн-кэш и Docker", "1.5 нед", "Н8-Н9"],
        ["17-18", "CI/CD", "0.5 нед", "Н9"],
        ["19-20", "Тестирование", "1.5 нед", "Н8-Н9"],
        ["21", "Документирование", "2 нед", "Н1-Н10"],
    ]
)

add_heading_styled("3.3 COCOMO — оценка трудоёмкости", 2)
add_para("Для оценки трудоёмкости использован метод COCOMO Intermediate (органический тип проекта).")
add_para("Исходные данные:", bold=True)
add_table(
    ["Параметр", "Значение"],
    [["Размер кода", "~9 000 LOC"],
     ["Backend (Java)", "~4 000 LOC"],
     ["Android (Kotlin)", "~4 000 LOC"],
     ["Тесты (Java/Kotlin)", "~1 000 LOC"]]
)

add_para("Расчёт (Basic COCOMO):", bold=True)
add_para("PM = 2.4 × (9.0)^1.05 = 2.4 × 10.2 ≈ 24.5 человеко-месяца")
add_para("TDEV = 2.5 × 24.5^0.38 ≈ 8.7 месяцев")
add_para("С учётом 15 факторов-драйверов (EAF = 0.37):")
add_para("PM_adj = 24.5 × 0.37 ≈ 9.1 человеко-месяца")
add_para("TDEV_adj = 2.5 × 9.1^0.38 ≈ 5.8 месяцев")
add_para("Результаты COCOMO:", bold=True)
add_table(
    ["Параметр", "Значение"],
    [["Трудоёмкость (Basic)", "24.5 чел.-мес."],
     ["Трудоёмкость (Intermediate)", "9.1 чел.-мес."],
     ["Время разработки", "~6 месяцев"],
     ["Средний размер команды", "2-3 разработчика"],
     ["Тип проекта", "Органический"]]
)

add_para("Расчёт ROI:", bold=True)
add_table(
    ["Параметр", "Значение"],
    [["Затраты на разработку", "1 219 600 ₽"],
     ["Прогноз годовых доходов", "2 100 000 ₽/год"],
     ["ROI (годовой)", "72.2%"],
     ["Срок окупаемости", "~7 месяцев"]]
)

doc.add_page_break()

# ============ ЗАКЛЮЧЕНИЕ ============
add_heading_styled("ЗАКЛЮЧЕНИЕ", 1)
add_para("В ходе выполнения курсовой работы разработана информационная система проката автомобилей «CarRentalApp», соответствующая требованиям методических указаний для траектории «Мобильная разработка».")

add_para("Основные результаты работы:", bold=True)
add_bullet("Проведён анализ предметной области, составлены BUC-диаграмма, SWOT-анализ и глоссарий")
add_bullet("Спроектирована архитектура PCMEF с интерфейсами IService/IRepository и ADR")
add_bullet("Разработана база данных PostgreSQL (6 таблиц, 3НФ, Flyway миграции)")
add_bullet("Реализован REST API (20+ endpoint'ов) с OpenAPI/Swagger документацией")
add_bullet("Разработано Android-приложение (7 экранов, Jetpack Compose, Material 3)")
add_bullet("Обеспечена безопасность (JWT, BCrypt, ролевая модель CLIENT/MANAGER/ADMIN)")
add_bullet("Реализованы паттерны GoF: State (жизненный цикл бронирования), Strategy (цены), Facade (бронирование)")
add_bullet("Обеспечен офлайн-режим через Room (кэш каталога и бронирований)")
add_bullet("Проведено тестирование (12 классов backend + Android, покрытие >40%)")
add_bullet("Настроена Docker-контейнеризация и CI/CD через GitHub Actions")
add_bullet("Выполнена экономическая оценка проекта (COCOMO: 9.1 чел.-мес., ROI: 72.2%)")

add_para("Заключение:", bold=True)
add_para("Разработанное приложение соответствует современным стандартам мобильной разработки, включает полный цикл бронирования автомобилей, обеспечивает безопасность данных и удобный интерфейс для всех категорий пользователей. Проект может быть использован как основа для коммерческого продукта после интеграции реального платёжного шлюза и системы push-уведомлений.")

doc.add_page_break()

# ============ СПИСОК ЛИТЕРАТУРЫ ============
add_heading_styled("СПИСОК ЛИТЕРАТУРЫ", 1)
references = [
    "1. Гагарина, Л.Г. Программная инженерия: учебное пособие / Л.Г. Гагарина. — М.: Форум, 2021. — 320 с.",
    "2. Орлов, С.А. Технологии разработки программного обеспечения / С.А. Орлов. — 10-е изд. — М.: Питер, 2018. — 672 с.",
    "3. Соммервилл, И. Инженерия программного обеспечения / И. Соммервилл. — М.: Вильямс, 2018. — 816 с.",
    "4. Гамма, Э. Приёмы объектно-ориентированного проектирования. Паттерны проектирования / Э. Гамма, Р. Хелм, Р. Джонсон, Дж. Влиссидес. — М.: Питер, 2016. — 368 с.",
    "5. Фаулер, М. Архитектура корпоративных программных приложений / М. Фаулер. — М.: Вильямс, 2018. — 544 с.",
    "6. Spring Boot Reference Documentation [Электронный ресурс]. — Режим доступа: https://docs.spring.io/spring-boot/docs/current/reference/htmlsingle/",
    "7. Android Developers Guide [Электронный ресурс]. — Режим доступа: https://developer.android.com/docs",
    "8. Jetpack Compose Documentation [Электронный ресурс]. — Режим доступа: https://developer.android.com/jetpack/compose/documentation",
    "9. PostgreSQL Documentation [Электронный ресурс]. — Режим доступа: https://www.postgresql.org/docs/",
    "10. Boehm, B. Software Engineering Economics / B. Boehm. — Prentice Hall, 1981. — 767 p.",
]
for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25)

# ============ СОХРАНЕНИЕ ============
output_path = "ПИЖ_Пояснительная_записка.docx"
doc.save(output_path)
print(f"OK: Document saved to {output_path}")
